# -*- coding: utf-8 -*-
"""Prueft ein Manuskript und sein PDF gegen alle Regeln, die je aufgefallen sind.

    python3 tools/buch-pruefen.py outputs/buch6/buch6-MANUSKRIPT.docx \
                                  outputs/buch6/buch6-Taschenbuch.pdf

Laeuft nach jeder Aenderung, nicht nur am Schluss. Wer nur prueft, was er
gerade geaendert hat, findet die Fehler von letzter Woche nie.
"""
import re
import sys

from docx import Document

try:
    import fitz
except ImportError:
    fitz = None

WORT = r'(?:der|die|das|dem|den|denen|deren|dessen)'
FRAGE = r'(?:was|wie|wer|wo|warum|wann|ob)'

# Woerter, nach denen die naechsten Muster keine Fehler sind
HARMLOS = {
    'so', 'genau', 'egal', 'nicht', 'auch', 'etwas', 'alles', 'nichts', 'mehr',
    'nur', 'immer', 'selbst', 'sogar', 'erst', 'noch', 'und', 'oder', 'aber',
    'denn', 'doch', 'dann', 'schon', 'eben', 'einfach', 'gerade', 'vielleicht',
    'ausser', 'wie', 'was', 'ist', 'sind', 'war', 'waren', 'wird', 'werden',
}

fehler = 0


def melde(name, treffer, zeigen=6):
    global fehler
    treffer = list(treffer)
    if treffer:
        fehler += len(treffer)
        print('PRUEF %s: %d' % (name, len(treffer)))
        for t in treffer[:zeigen]:
            print('       %s' % (t,))
        if len(treffer) > zeigen:
            print('       ... und %d weitere' % (len(treffer) - zeigen))
    else:
        print('OK    %s' % name)


def zeilen(docx):
    for p in docx.paragraphs:
        for z in p.text.split('\n'):
            # Word-Inhaltsverzeichnis ueberspringen, das ist eine Kopie und
            # wird beim Oeffnen ohnehin neu aufgebaut
            if z.strip() and '\t' not in z:
                yield z.strip()


def kommas(docx):
    """Fehlende Kommas vor Neben- und Relativsaetzen."""
    muster = [
        r'\b([A-ZÄÖÜ]\w+)\s+(' + WORT + r')\s+\w+\s+\w+',
        r'\b(\w{3,})\s+(an|in|mit|von|zu|auf|für|bei|nach|aus)\s+(dem|der|denen)\s+\w+\s+\w+',
        r'\b(\w{3,})\s+(' + FRAGE + r')\s+(ich|du|er|sie|es|wir|man|jemand|niemand)\b',
        r'\b(\w{3,})\s+(dass|weil|obwohl|damit|während|bevor|nachdem|falls|sobald|bis|ohne dass)\s+\w+',
        r'\b(\w{3,})\s+als\s+(ich|du|er|sie|es|wir|man|alles|wären|würde)\b',
    ]
    for z in zeilen(docx):
        for m in muster:
            for tref in re.finditer(m, z):
                vor = z[:tref.start()].rstrip()
                if not vor or vor.endswith((',', '.', '!', '?', ':', ';', '„', '“', '-')):
                    continue
                if tref.group(1).lower() in HARMLOS:
                    continue
                yield z[max(0, tref.start() - 40):tref.end() + 15]
                break


def pruefe_docx(pfad):
    d = Document(pfad)
    ps = [p.text for p in d.paragraphs]
    T = '\n'.join(ps)

    print('--- Text: %s' % pfad)
    melde('scharfes s', re.findall(r'.{15}ß.{15}', T))
    melde('Gedankenstriche', re.findall(r'.{15}[–—].{15}', T))
    melde('gerade Anführungszeichen', [z[:60] for z in ps if '"' in z])
    melde('gerader Apostroph', [z[:60] for z in ps if "'" in z])
    melde('doppelte Leerzeichen',
          [z[:60] for z in ps if '  ' in z.strip() and '·' not in z])
    melde('Leerzeichen vor Satzzeichen', re.findall(r'.{15} [,.;:!?].{8}', T))
    melde('Auslassungspunkte', [z[:60] for z in ps if '...' in z or '…' in z])

    auf, zu = T.count('„'), T.count('“')
    melde('Anführungszeichen paarig',
          [] if auf == zu else ['%d auf, %d zu' % (auf, zu)])

    melde('fehlende Kommas', kommas(d))

    # Rechtssicherheit, siehe context/rechtssicherheit.md
    print('--- Rechtssicherheit')
    melde('private Mailadresse', [z for z in ps if 'bluewin' in z])
    melde('Bezeichnung Heilerin', [z for z in ps if 'Heiler' in z])
    melde('Therapeutin über die Autorin',
          [z for z in ps if 'Therapeut' in z
           and 'Psychotherapeut' not in z and 'Freundin' not in z])
    melde('Ausschlussdiagnose',
          [z for z in ps if re.search(r'\b(Das|Es) ist kein (Burnout|Depression)', z)])
    melde('Heilversprechen',
          [z[:70] for z in ps
           if re.search(r'\b(heilt dich|macht dich gesund|befreit dich|kuriert)\b', z)])

    for pflicht in ('Haftungsausschluss', 'Zur Autorin', 'Notfall und professionelle Hilfe',
                    'Fallbeispiele und Anonymisierung', 'Urheberrecht', 'Markenrecht',
                    'beyondlimitsnow25@gmail.com', 'Wiesentalstrasse 68'):
        melde('Pflichtteil: %s' % pflicht,
              [] if any(pflicht in z for z in ps) else ['fehlt im Text'])

    # Kapitelnummern lueckenlos
    kap = sorted({int(m.group(1)) for z in ps
                  for m in [re.match(r'^Kapitel (\d+):', z.strip())] if m})
    melde('Kapitelnummern lückenlos',
          [] if kap == list(range(1, len(kap) + 1)) else ['gefunden: %s' % kap])
    melde('Kapiteltitel mit Schlusspunkt',
          [z.strip() for z in ps
           if re.match(r'^Kapitel \d+:', z.strip()) and z.strip().endswith('.')])


def pruefe_pdf(pfad):
    if fitz is None:
        print('PRUEF PyMuPDF fehlt, PDF nicht geprüft')
        return
    doc = fitz.open(pfad)
    print('--- Satz: %s, %d Seiten' % (pfad, len(doc)))

    melde('leere Seiten',
          ['Seite %d' % (i + 1) for i, s in enumerate(doc)
           if len(s.get_text().strip()) < 6])
    melde('fast leere Seiten',
          ['Seite %d, nur %d Zeichen' % (i + 1, len(s.get_text().strip()))
           for i, s in enumerate(doc)
           if 6 <= len(s.get_text().strip()) < 100 and i > 2])

    # KDP lehnt PDFs ab, in denen eine Schrift nicht eingebettet ist
    schriften = set()
    for s in doc:
        schriften |= {(f[3], f[1]) for f in s.get_fonts()}
    melde('Schriften eingebettet',
          ['%s ist nur referenziert' % n for n, art in schriften if art == 'n/a'])

    # Seitenmass muss zum bestellten Format passen
    r = doc[0].rect
    mass = '%.1f x %.1f mm' % (r.width / 72 * 25.4, r.height / 72 * 25.4)
    melde('Seitenmass A5 (148,0 x 210,0 mm)',
          [] if mass == '148.0 x 210.0 mm' else [mass])

    # Inhaltsverzeichnis gegen die echten Seitenzahlen
    zei = []
    for i in range(min(6, len(doc))):
        zei += [z.strip() for z in doc[i].get_text().split('\n') if z.strip()]
    eintraege = []
    for k, z in enumerate(zei):
        m = re.match(r'^[\. ]*(\d+)$', z)
        if m and k + 1 < len(zei):
            eintraege.append((zei[k + 1], int(m.group(1))))
    seiten = [re.sub(r'\s+', ' ', s.get_text().replace('\n', ' ')) for s in doc]
    falsch = []
    for titel, nr in eintraege:
        schl = re.sub(r'\s+', ' ', titel)[:40]
        tref = [i + 1 for i, s in enumerate(seiten) if schl in s and i >= 4]
        if not tref:
            falsch.append('%s: im Text nicht gefunden' % titel[:45])
        elif tref[0] != nr:
            falsch.append('%s: Verzeichnis %d, tatsächlich %d' % (titel[:40], nr, tref[0]))
    melde('Inhaltsverzeichnis (%d Einträge)' % len(eintraege), falsch)

    # Krisennummern muessen weit vorne stehen, nicht nur im Rechtlichen
    vorne = [i + 1 for i, s in enumerate(doc) if '143' in s.get_text() and i < len(doc) // 3]
    melde('Krisennummern im vorderen Drittel',
          [] if vorne else ['nur hinten im Rechtlichen'])

    ruecken = len(doc) * 0.0025 * 25.4
    print('      Umfang %d Seiten, Rücken %.2f mm, Cover %.2f x 216,40 mm'
          % (len(doc), ruecken, 3.2 + 148 + ruecken + 148 + 3.2))


if __name__ == '__main__':
    if len(sys.argv) < 2:
        raise SystemExit(__doc__)
    pruefe_docx(sys.argv[1])
    if len(sys.argv) > 2:
        pruefe_pdf(sys.argv[2])
    print()
    print('Auffälligkeiten insgesamt: %d' % fehler)
    print('Nicht jede ist ein Fehler. Jede muss angeschaut werden.')
    sys.exit(1 if fehler else 0)
